import docx

def parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")
