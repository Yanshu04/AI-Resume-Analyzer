import os
import docx
from app.models.schemas import ParsedResume

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def generate_markdown_content(resume: ParsedResume) -> str:
    """
    Format parsed resume into clean, structured Markdown text.
    """
    links_str = " | ".join(resume.links)
    md = []
    md.append(f"# {resume.name.upper()}")
    md.append(f"**Email:** {resume.email} | **Phone:** {resume.phone}")
    if links_str:
        md.append(f"**Links:** {links_str}")
    md.append("\n---\n")
    
    # Skills
    md.append("## TECHNICAL SKILLS")
    md.append(", ".join(resume.skills))
    md.append("")
    
    # Experience
    md.append("## PROFESSIONAL EXPERIENCE")
    for exp in resume.experience:
        md.append(f"### {exp.title} — {exp.company}")
        md.append(f"*{exp.duration}*")
        for bullet in exp.bullets:
            md.append(f"- {bullet}")
        md.append("")
        
    # Projects
    if resume.projects:
        md.append("## TECHNICAL PROJECTS")
        for proj in resume.projects:
            tech_str = f" (*{', '.join(proj.tech)}*)" if proj.tech else ""
            md.append(f"### {proj.name}{tech_str}")
            md.append(proj.description)
            md.append("")
            
    # Education
    md.append("## EDUCATION")
    for edu in resume.education:
        md.append(f"### {edu.degree}")
        md.append(f"{edu.institution} | *Graduation Year: {edu.year}*")
        md.append("")
        
    # Certifications
    if resume.certifications:
        md.append("## CERTIFICATIONS")
        for cert in resume.certifications:
            md.append(f"- {cert}")
            
    return "\n".join(md)

def export_to_markdown(resume: ParsedResume, file_path: str) -> None:
    """
    Saves the resume as a Markdown file.
    """
    content = generate_markdown_content(resume)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

def export_to_docx(resume: ParsedResume, file_path: str) -> None:
    """
    Saves the resume as a Microsoft Word Document (DOCX).
    """
    doc = docx.Document()
    
    # Page settings: 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(0.75)
        section.bottom_margin = docx.shared.Inches(0.75)
        section.left_margin = docx.shared.Inches(0.75)
        section.right_margin = docx.shared.Inches(0.75)

    # 1. Header (Name, Contacts)
    title_p = doc.add_paragraph()
    title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(resume.name.upper())
    title_run.font.name = 'Arial'
    title_run.font.size = docx.shared.Pt(18)
    title_run.bold = True
    
    contacts_p = doc.add_paragraph()
    contacts_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    links_str = f" | { ' | '.join(resume.links) }" if resume.links else ""
    contacts_run = contacts_p.add_run(f"{resume.email} | {resume.phone}{links_str}")
    contacts_run.font.name = 'Arial'
    contacts_run.font.size = docx.shared.Pt(9.5)
    
    # 2. Skills
    doc.add_heading("TECHNICAL SKILLS", level=1)
    skills_p = doc.add_paragraph(", ".join(resume.skills))
    skills_p.style.font.name = 'Arial'
    
    # 3. Experience
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    for exp in resume.experience:
        exp_p = doc.add_paragraph()
        title_run = exp_p.add_run(f"{exp.title} at {exp.company}")
        title_run.bold = True
        title_run.font.size = docx.shared.Pt(11)
        
        dur_run = exp_p.add_run(f"\t{exp.duration}")
        dur_run.font.size = docx.shared.Pt(10)
        dur_run.italic = True
        
        # Bullets
        for bullet in exp.bullets:
            doc.add_paragraph(bullet, style='List Bullet')
            
    # 4. Projects
    if resume.projects:
        doc.add_heading("TECHNICAL PROJECTS", level=1)
        for proj in resume.projects:
            proj_p = doc.add_paragraph()
            name_run = proj_p.add_run(proj.name)
            name_run.bold = True
            
            if proj.tech:
                tech_run = proj_p.add_run(f" ({', '.join(proj.tech)})")
                tech_run.font.size = docx.shared.Pt(9.5)
                tech_run.italic = True
                
            doc.add_paragraph(proj.description)
            
    # 5. Education
    doc.add_heading("EDUCATION", level=1)
    for edu in resume.education:
        edu_p = doc.add_paragraph()
        edu_run = edu_p.add_run(edu.degree)
        edu_run.bold = True
        
        inst_run = edu_p.add_run(f"\t{edu.institution} ({edu.year})")
        inst_run.italic = True
        
    # 6. Certifications
    if resume.certifications:
        doc.add_heading("CERTIFICATIONS", level=1)
        for cert in resume.certifications:
            doc.add_paragraph(cert, style='List Bullet')
            
    # Apply standard fonts to headings
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading'):
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.color.rgb = docx.shared.RGBColor(13, 13, 13)
                
    doc.save(file_path)

def export_to_pdf(resume: ParsedResume, file_path: str) -> None:
    """
    Saves the resume as a formatted PDF Document (using ReportLab).
    """
    # Setup document template with 0.5-inch margins
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Define custom styles
    name_style = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1, # Center
        textColor=colors.HexColor('#0D0D0D'),
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'ResumeContacts',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        alignment=1, # Center
        textColor=colors.HexColor('#4A4A4A'),
        spaceAfter=10
    )
    
    h1_style = ParagraphStyle(
        'ResumeH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#0D0D0D'),
        spaceBefore=8,
        spaceAfter=2,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#222222'),
        spaceAfter=4
    )
    
    bold_sub_style = ParagraphStyle(
        'ResumeSubBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#1A1A1A'),
        keepWithNext=True
    )
    
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        leftIndent=15,
        firstLineIndent=-10,
        textColor=colors.HexColor('#2E2E2E'),
        spaceAfter=2
    )

    story = []
    
    # 1. Header (Name, Contacts)
    story.append(Paragraph(resume.name.upper(), name_style))
    links_str = f" | { ' | '.join(resume.links) }" if resume.links else ""
    contacts = f"{resume.email}  |  {resume.phone}{links_str}"
    story.append(Paragraph(contacts, contact_style))
    
    # Helper to append a horizontal line
    def add_section_header(title: str):
        story.append(Paragraph(title, h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2A2A2A"), spaceBefore=1, spaceAfter=6))

    # 2. Skills
    add_section_header("TECHNICAL SKILLS")
    story.append(Paragraph(", ".join(resume.skills), body_style))
    story.append(Spacer(1, 6))

    # 3. Experience
    add_section_header("PROFESSIONAL EXPERIENCE")
    for exp in resume.experience:
        # Title & Duration in a single paragraph flowable
        title_p = f"<b>{exp.title}</b> at <i>{exp.company}</i>"
        dur_p = f"<font color='#555555'>{exp.duration}</font>"
        # Format left-right align table or simple paragraph
        story.append(Paragraph(f"{title_p} &nbsp;&bull;&nbsp; {dur_p}", bold_sub_style))
        
        # Bullets
        for bullet in exp.bullets:
            story.append(Paragraph(f"&bull;&nbsp; {bullet}", bullet_style))
        story.append(Spacer(1, 4))
        
    # 4. Projects
    if resume.projects:
        add_section_header("TECHNICAL PROJECTS")
        for proj in resume.projects:
            tech_str = f" ({', '.join(proj.tech)})" if proj.tech else ""
            story.append(Paragraph(f"<b>{proj.name}</b><i>{tech_str}</i>", bold_sub_style))
            story.append(Paragraph(proj.description, body_style))
            story.append(Spacer(1, 4))

    # 5. Education
    add_section_header("EDUCATION")
    for edu in resume.education:
        story.append(Paragraph(f"<b>{edu.degree}</b> &nbsp;&bull;&nbsp; <i>{edu.institution}</i> ({edu.year})", body_style))
        story.append(Spacer(1, 3))
        
    # 6. Certifications
    if resume.certifications:
        add_section_header("CERTIFICATIONS")
        for cert in resume.certifications:
            story.append(Paragraph(f"&bull;&nbsp; {cert}", bullet_style))
            
    doc.build(story)
